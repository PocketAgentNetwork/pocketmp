#!/usr/bin/env python3
"""
Resume Generator Script
Converts structured resume data to PDF or DOCX format
"""

import json
import sys
from pathlib import Path

def generate_docx(data, output_path):
    """Generate DOCX resume (ATS-friendly)"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Header
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = header.add_run(data['header']['name'])
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{data['header']['location']} | {data['header']['phone']} | {data['header']['email']}"
    if data['header'].get('linkedin'):
        contact_text += f" | {data['header']['linkedin']}"
    if data['header'].get('portfolio'):
        contact_text += f" | {data['header']['portfolio']}"
    contact.add_run(contact_text).font.size = Pt(10)
    
    if data['header'].get('target_role'):
        target = doc.add_paragraph()
        target.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target_run = target.add_run(f"TARGET ROLE: {data['header']['target_role']}")
        target_run.font.size = Pt(10)
    
    doc.add_paragraph()  # Spacing
    
    # Professional Summary
    if data.get('summary'):
        summary_heading = doc.add_paragraph()
        summary_heading.add_run('PROFESSIONAL SUMMARY').bold = True
        summary_heading.add_run().add_break()
        doc.add_paragraph(data['summary'])
        doc.add_paragraph()  # Spacing
    
    # Core Competencies
    if data.get('skills'):
        skills_heading = doc.add_paragraph()
        skills_heading.add_run('CORE COMPETENCIES').bold = True
        
        for category, items in data['skills'].items():
            skill_para = doc.add_paragraph()
            skill_para.add_run(f"{category.upper()}: ").bold = True
            skill_para.add_run(' • '.join(items))
        
        doc.add_paragraph()  # Spacing
    
    # Professional Experience
    if data.get('experience'):
        exp_heading = doc.add_paragraph()
        exp_heading.add_run('PROFESSIONAL EXPERIENCE').bold = True
        
        for job in data['experience']:
            # Company and dates
            job_header = doc.add_paragraph()
            company_run = job_header.add_run(f"{job['company']} | {job['location']}")
            company_run.bold = True
            
            # Title and dates on same line
            title_para = doc.add_paragraph()
            title_para.add_run(job['title']).italic = True
            title_para.add_run(f" | {job['start_date']} – {job['end_date']}")
            
            # Bullets
            for bullet in job['bullets']:
                bullet_para = doc.add_paragraph(bullet, style='List Bullet')
                bullet_para.paragraph_format.left_indent = Inches(0.25)
            
            doc.add_paragraph()  # Spacing between jobs
    
    # Education
    if data.get('education'):
        edu_heading = doc.add_paragraph()
        edu_heading.add_run('EDUCATION').bold = True
        
        for edu in data['education']:
            edu_para = doc.add_paragraph()
            edu_para.add_run(f"{edu['degree']}, {edu['major']}").bold = True
            edu_para.add_run(f"\n{edu['university']}, {edu['location']} | {edu['graduation_year']}")
        
        doc.add_paragraph()  # Spacing
    
    # Certifications
    if data.get('certifications'):
        cert_heading = doc.add_paragraph()
        cert_heading.add_run('CERTIFICATIONS').bold = True
        
        for cert in data['certifications']:
            cert_para = doc.add_paragraph(style='List Bullet')
            cert_text = f"{cert['name']}, {cert['issuer']}, {cert['year']}"
            if cert.get('expiration'):
                cert_text += f" – Expires {cert['expiration']}"
            cert_para.add_run(cert_text)
    
    # Save
    doc.save(output_path)
    print(f"✅ Resume saved to: {output_path}")


def generate_pdf(data, output_path):
    """Generate PDF resume"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        print("Error: reportlab not installed. Run: pip install reportlab")
        sys.exit(1)
    
    doc = SimpleDocTemplate(output_path, pagesize=letter,
                           topMargin=0.75*inch, bottomMargin=0.75*inch,
                           leftMargin=0.75*inch, rightMargin=0.75*inch)
    
    styles = getSampleStyleSheet()
    story = []
    
    # Custom styles
    name_style = ParagraphStyle('Name', parent=styles['Heading1'],
                                fontSize=16, alignment=TA_CENTER, spaceAfter=6)
    contact_style = ParagraphStyle('Contact', parent=styles['Normal'],
                                   fontSize=10, alignment=TA_CENTER, spaceAfter=12)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'],
                                   fontSize=12, spaceAfter=6, spaceBefore=12)
    
    # Header
    story.append(Paragraph(data['header']['name'], name_style))
    
    contact_text = f"{data['header']['location']} | {data['header']['phone']} | {data['header']['email']}"
    if data['header'].get('linkedin'):
        contact_text += f" | {data['header']['linkedin']}"
    if data['header'].get('portfolio'):
        contact_text += f" | {data['header']['portfolio']}"
    story.append(Paragraph(contact_text, contact_style))
    
    if data['header'].get('target_role'):
        story.append(Paragraph(f"<b>TARGET ROLE:</b> {data['header']['target_role']}", contact_style))
    
    story.append(Spacer(1, 0.2*inch))
    
    # Professional Summary
    if data.get('summary'):
        story.append(Paragraph('<b>PROFESSIONAL SUMMARY</b>', heading_style))
        story.append(Paragraph(data['summary'], styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    # Core Competencies
    if data.get('skills'):
        story.append(Paragraph('<b>CORE COMPETENCIES</b>', heading_style))
        for category, items in data['skills'].items():
            skill_text = f"<b>{category.upper()}:</b> {' • '.join(items)}"
            story.append(Paragraph(skill_text, styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    # Professional Experience
    if data.get('experience'):
        story.append(Paragraph('<b>PROFESSIONAL EXPERIENCE</b>', heading_style))
        
        for job in data['experience']:
            # Company
            story.append(Paragraph(f"<b>{job['company']} | {job['location']}</b>", styles['Normal']))
            # Title and dates
            story.append(Paragraph(f"<i>{job['title']}</i> | {job['start_date']} – {job['end_date']}", styles['Normal']))
            
            # Bullets
            bullet_items = [ListItem(Paragraph(bullet, styles['Normal']), leftIndent=20) 
                          for bullet in job['bullets']]
            story.append(ListFlowable(bullet_items, bulletType='bullet'))
            story.append(Spacer(1, 0.1*inch))
    
    # Education
    if data.get('education'):
        story.append(Paragraph('<b>EDUCATION</b>', heading_style))
        for edu in data['education']:
            story.append(Paragraph(f"<b>{edu['degree']}, {edu['major']}</b>", styles['Normal']))
            story.append(Paragraph(f"{edu['university']}, {edu['location']} | {edu['graduation_year']}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
    
    # Certifications
    if data.get('certifications'):
        story.append(Paragraph('<b>CERTIFICATIONS</b>', heading_style))
        for cert in data['certifications']:
            cert_text = f"{cert['name']}, {cert['issuer']}, {cert['year']}"
            if cert.get('expiration'):
                cert_text += f" – Expires {cert['expiration']}"
            story.append(Paragraph(f"• {cert_text}", styles['Normal']))
    
    # Build PDF
    doc.build(story)
    print(f"✅ Resume saved to: {output_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python generate_resume.py <input.json> <output.pdf|output.docx>")
        print("\nExample:")
        print("  python generate_resume.py resume_data.json resume.docx")
        print("  python generate_resume.py resume_data.json resume.pdf")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    # Load data
    try:
        with open(input_file, 'r') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Error: Input file '{input_file}' not found")
        sys.exit(1)
    except json.JSONDecodeError:
        print(f"Error: Invalid JSON in '{input_file}'")
        sys.exit(1)
    
    # Generate based on extension
    output_path = Path(output_file)
    if output_path.suffix.lower() == '.docx':
        generate_docx(data, output_file)
    elif output_path.suffix.lower() == '.pdf':
        generate_pdf(data, output_file)
    else:
        print("Error: Output file must be .pdf or .docx")
        sys.exit(1)


if __name__ == "__main__":
    main()
